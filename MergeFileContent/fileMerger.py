import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import json
import PyPDF2
from docx import Document

def merge_files():
    """Open file dialog to select multiple files and merge them intelligently."""
    root = tk.Tk()
    root.withdraw()  # Hide main Tkinter window

    # Select multiple files
    files = filedialog.askopenfilenames(title="Select files to merge",
                                        filetypes=[("All Supported Files", "*.txt;*.csv;*.json;*.pdf;*.docx"),
                                                   ("Text files", "*.txt"),
                                                   ("CSV files", "*.csv"),
                                                   ("JSON files", "*.json"),
                                                   ("PDF files", "*.pdf"),
                                                   ("Word files", "*.docx")])

    if not files:
        messagebox.showwarning("No Files Selected", "Please select at least one file.")
        return

    # Ask user where to save the merged file
    save_path = filedialog.asksaveasfilename(title="Save Merged File As",
                                             filetypes=[("Text file", "*.txt"),
                                                        ("CSV file", "*.csv"),
                                                        ("JSON file", "*.json"),
                                                        ("PDF file", "*.pdf"),
                                                        ("Word file", "*.docx")])

    if not save_path:
        return  # User canceled save

    try:
        if save_path.endswith(".txt"):
            with open(save_path, "w", encoding="utf-8") as outfile:
                for file in files:
                    with open(file, "r", encoding="utf-8") as infile:
                        outfile.write(infile.read() + "\n")
        
        elif save_path.endswith(".csv"):
            df_list = [pd.read_csv(file) for file in files if file.endswith(".csv")]
            merged_df = pd.concat(df_list, ignore_index=True)
            merged_df.to_csv(save_path, index=False)
        
        elif save_path.endswith(".json"):
            json_data = []
            for file in files:
                with open(file, "r", encoding="utf-8") as infile:
                    json_data.append(json.load(infile))
            with open(save_path, "w", encoding="utf-8") as outfile:
                json.dump(json_data, outfile, indent=4)
        
        elif save_path.endswith(".pdf"):
            merger = PyPDF2.PdfMerger()
            for file in files:
                if file.endswith(".pdf"):
                    merger.append(file)
            merger.write(save_path)
            merger.close()
        
        elif save_path.endswith(".docx"):
            merged_doc = Document()
            for file in files:
                if file.endswith(".docx"):
                    doc = Document(file)
                    for para in doc.paragraphs:
                        merged_doc.add_paragraph(para.text)
            merged_doc.save(save_path)
        
        messagebox.showinfo("Success", f"Files merged successfully into:\n{save_path}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{str(e)}")

if __name__ == "__main__":
    merge_files()